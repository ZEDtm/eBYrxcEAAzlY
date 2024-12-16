import logging
import os

import aiohttp
import requests
import gspread
import pandas as pd
import requests
from google.auth.transport.requests import Request
from google.oauth2.service_account import Credentials
from datetime import datetime
import pytz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from dateutil.relativedelta import relativedelta
import calendar
import subprocess
from admin_panel.models import TgUser, Admin
from BotInvest import settings
import convertapi
name_months = {
    1: 'января',
    2: 'февраля',
    3: 'марта',
    4: 'апреля',
    5: 'мая',
    6: 'июня',
    7: 'июля',
    8: 'августа',
    9: 'сентября',
    10: 'октября',
    11: 'ноября',
    12: 'декабря'
}
name_months_i = {
    1: 'январь',
    2: 'февраль',
    3: 'март',
    4: 'апрель',
    5: 'май',
    6: 'июнь',
    7: 'июль',
    8: 'август',
    9: 'сентябрь',
    10: 'октябрь',
    11: 'ноябрь',
    12: 'декабрь'
}


def download_sheet():
    excel_path = 'admin_panel/report_data/data.xlsx'
    # Установите область доступа
    scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
    # Загрузите учетные данные
    creds = Credentials.from_service_account_file('admin_panel/report_data/api_key_excel.json', scopes=scope)
    # Обновите токен доступа
    creds.refresh(Request())

    client = gspread.authorize(creds)

    # Откройте таблицу по имени
    spreadsheet = client.open('Отчет Агента')

    # Получите ID таблицы
    spreadsheet_id = spreadsheet.id

    # URL для скачивания в формате Excel
    url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=xlsx"
    # Выполните GET-запрос для скачивания файла
    response = requests.get(url, headers={'Authorization': f'Bearer {creds.token}'})

    # Сохраните файл

    with open(excel_path, "wb") as f:
        f.write(response.content)
    get_data(excel_path)


def get_data(excel_path: str):
    data = pd.read_excel(excel_path, 'Октябрь', usecols='A:Q')
    for i in range(1, data.shape[0]):

        # Извлечение данных с EXCEL
        fio = data['Unnamed: 0'].iloc[i]
        if pd.isna(fio):
            break
        num_dogovora = data['Unnamed: 1'].iloc[i]
        date_ot = data['Unnamed: 2'].iloc[i]
        ser_num = data['Unnamed: 3'].iloc[i]
        sum_1 = data['Unnamed: 13'].iloc[i]
        sum_2 = data['Unnamed: 14'].iloc[i]
        sum_3 = data['Unnamed: 15'].iloc[i]
        itogo = data['Unnamed: 16'].iloc[i]
        create_docx(fio, num_dogovora, date_ot, ser_num, sum_1, sum_2, sum_3, itogo)
    send_users()


def create_docx(fio: str, num_dogovora: str, date_ot:str, ser_num:str, sum_1:str, sum_2:str, sum_3:str, itogo:str):
    last_day, last_month, current_date = get_past_month()

    doc = Document()
    # Получаем доступ к секции документа
    section = doc.sections[0]

    # Устанавливаем ориентацию страницы на альбомную
    section.orientation = 1  # 1 - альбомная, 0 - портретная

    # Устанавливаем размеры страницы (по умолчанию A4)
    section.page_width = Inches(11)  # Ширина для альбомной ориентации
    section.page_height = Inches(8.5)  # Высота для альбомной ориентации

    paragrahp = doc.add_paragraph('Отчет агента')
    paragrahp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrahp = doc.add_paragraph(f'по агентскому договору №{num_dogovora} от {date_ot}.')
    paragrahp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    doc.add_paragraph(
        f"г. Москва                                                                                                                                                                                                           «{current_date.day}» {name_months[current_date.month]} {current_date.year} г.  ")
    doc.add_paragraph()
    doc.add_paragraph(
        f'Индивидуальный предприниматель Холин Сергей Сергеевич именуемый далее «Агент», направляет настоящий отчет об исполнении поручения ')
    doc.add_paragraph(
        f'Принципалу {fio} в рамках агентского договора №{num_dogovora} от {date_ot}')
    doc.add_paragraph()
    doc.add_paragraph(
        f'Отчет составлен за отчетный период с «1» {name_months[last_month.month]} {last_month.year} г. по «{last_day}» {name_months[last_month.month]} {last_month.year} г. (далее – «ОП»).')
    doc.add_paragraph(
        'Оборудование Принципала, в отношении которого было выполнено агентское поручение и в отношении которого предоставляется настоящий отчет: ')
    doc.add_paragraph(f'{ser_num}  (далее – «Оборудование») ')
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)

    row1 = table.rows[0]
    row2 = table.rows[1]
    row3 = table.rows[2]

    row1.cells[0].text = '1. Общая сумма расходов на электричество за отчетный период: '
    row1.cells[1].text = str(sum_1) + ' руб.'
    row2.cells[0].text = '2. Общая сумма иных расходов за отчетный период (с конкретизацией): '
    row2.cells[1].text = str(sum_2) + ' руб.'
    row3.cells[0].text = '3. Общая сумма агентского вознаграждения за отчетный период: '
    row3.cells[1].text = str(sum_3) + ' руб.'
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph(
        f'Итого общая сумма денежных средств, которую Принципал обязан перечислить Агенту за ОП: {itogo} руб.')
    doc.add_paragraph()
    doc.add_paragraph(
        'Принципал с отчетом ознакомился, претензий к его содержанию, а также претензий к выполненному Агентом агентскому поручению за отчетный период ')
    doc.add_paragraph('Принципал не имеет. ')
    doc.add_paragraph()
    doc.add_paragraph(
        'В случае, если в течение 5 дней с момента получения настоящего Отчета Принципалом не будут представлены на него возражения, Отчет считается принятым и согласованным Принципалом в полном объеме.')
    doc.add_paragraph()
    doc.add_paragraph(f'Агент: ___________ / Холин С. С')
    doc.add_paragraph()
    doc.add_paragraph(f'Принципал: ___________ / {short_fio(fio)}')

    for paragraph in doc.paragraphs:
        paragraph.style = doc.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
        p_frm = paragraph.paragraph_format
        p_frm.space_after = Pt(0)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.3)
        section.right_margin = Cm(1.3)

    doc.save(f'admin_panel/report_data/users/{fio}.docx')
    generate_pdf(f'admin_panel/report_data/users/{fio}.docx',f'admin_panel/report_data/users/{fio}.pdf')

def get_past_month():
    moscow_tz = pytz.timezone('Europe/Moscow')

    current_date = datetime.now(tz=moscow_tz)

    # Вычисляем прошлый месяц
    last_month = current_date - relativedelta(months=1)
    last_day = calendar.monthrange(last_month.year, last_month.month)[1]
    return last_day, last_month, current_date


def short_fio(fio:str):
    f = fio.strip().split(' ')[0] # Фамилия
    n = fio.strip().split(' ')[1][0] # Имя
    o = fio.strip().split(' ')[2][0] # Отчество
    return f'{f} {n}. {o}.'




def generate_pdf(doc_path, path):
    convertapi.api_secret = settings.SECRET_API
    convertapi.convert('pdf', {
        'File': doc_path
    }, from_format='docx').save_files(path)
    os.remove(doc_path)


def send_users():
    pdf_files = os.listdir(f'admin_panel/report_data/users')
    for file in pdf_files:
        user_id = get_telegram_id_by_fio(file.replace('.pdf', ''))
        if user_id is None:
            users = Admin.objects.all()
            for user in users:
                send_file_to_telegram(file.replace('.pdf', ''), int(user.telegram_id))
        response = send_file_to_telegram(file.replace('.pdf', ''), int(user_id))
        if response:
            data = {
                "id": response['result']['message_id'],  # message_id из ответа телеги
                "text": response['result']['caption'],  # message_text
                "status": "sent",  # по умолчанию
                "sender": 0,  # по умолчанию
                "recipient": response['result']['chat']['id'],  # chat_id
                "date": response['result']['date'],  # date из ответа телеги
                "files":
                    [{
                        "file_name": file,  # рандомное имя
                        "file_id": response['result']['document']['file_id'],  # из ответа телеги
                        "file_size": response['result']['document']['file_size'],  # из ответа телеги
                        "mime_type": response['result']['document']['mime_type']  # из ответа телеги
                    }]
            }
            headers = {
                'Authorization': f'Bearer {settings.TG_TOKEN_BOT}'
            }
            async with aiohttp.ClientSession() as session:
                async with session.post(settings.FASTAPI_URL, headers=headers, json=data) as response:
                    if response.status != 200:
                        logging.error(f"Failed to send message to FastAPI: {response.text()}")
                    else:
                        logging.info("Message sent to FastAPI successfully")
        else:
            users = Admin.objects.all()
            for user in users:
                send_file_to_telegram(file.replace('.pdf', ''), int(user.telegram_id))








def get_telegram_id_by_fio(fio):
    try:
        user = TgUser.objects.get(fio=fio)
        return user.telegram_id
    except TgUser.DoesNotExist:
        return None  # или обработайте случай, когда пользователь не найден


def send_file_to_telegram(fio, chat_id):
    last_month = get_past_month()[1]
    url = f"https://api.telegram.org/bot{settings.TG_TOKEN_BOT}/sendDocument"

    with open(f'admin_panel/report_data/users/{fio}.pdf', 'rb') as file:
        files = {'document': file}
        data = {
            'chat_id': chat_id,
            'caption': 'Уважаемый инвестор 🤝 \n\n'
                       f'Согласно нашему агентскому договору, высылаем Вам отчет агента за {name_months_i[last_month.month]}.\n\n'
                       f'<i>Данный документ считается принятым, если с вашей стороны в течении 5 дней после получения данного сообщения, не будут направлены мотивированные возражения.</i>\n\n'
                       f'<u>Документ оплачивать не нужно</u>',
            'parse_mode': 'HTML'# Текстовое сообщение, прикрепленное к файлу
        }

        response = requests.post(url, data=data, files=files)

        if response.status_code == 200:
            return response.json()
        else:
            return False
if __name__ == '__main__':
    download_sheet(excel_path="")





