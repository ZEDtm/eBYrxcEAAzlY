import datetime

import convertapi
from django.conf import settings
from docx import Document
from loguru import logger

from admin_panel.models import TgUser
from bot.get_now_month import month_name_to_number
from bot.service_async import get_user_workers


async def create_pdf_uniq_code(tg_user: TgUser):
    workers = await get_user_workers(tg_user)
    list_workers = ''
    for work in workers:
        list_workers += f'{work},\n'

    date_now = datetime.datetime.now()
    name, _ = month_name_to_number()

    replace_dict = generate_replace_word(
        telegram_id=tg_user.telegram_id,
        day=date_now.day,
        month=f'{name}',
        name=f'{tg_user.fio}',
        seria=f'{tg_user.passport[:4]}',
        number=f'{tg_user.passport[4:]}',
        kem_vidan=f'{tg_user.get_passport}',
        date_passport=f'{tg_user.date_passport}',
        address=f'{tg_user.address_registration}',
    )

    create_word_file(
        replace_dict=replace_dict,
        list_work=list_workers,
        len_worker=f'{len(workers)}',
        name=f'{tg_user.fio}',
        seria=f'{tg_user.passport[:4]}',
        number=f'{tg_user.passport[4:]}',
        kem_vidan=f'{tg_user.get_passport}',
        date_passport=f'{tg_user.date_passport}',
        address=f'{tg_user.address_registration}',
        account_number=f'{tg_user.number_account}',
        bank_name=f'{tg_user.bank_name}',
        bic=f'{tg_user.bank_bik}',
        corr=f'{tg_user.corr_account}',
        telegram_id=tg_user.telegram_id,
    )




def generate_replace_word(
        telegram_id: int,
        day: int,
        month: str,
        name: str,
        seria: str,
        number: str,
        kem_vidan: str,
        date_passport: str,
        address: str,
) -> dict:
    return {
        'Агентский договор № _______': f'Агентский договор № {telegram_id}',
        'г.  Москва.    		           							 «13» июля 2024 г.  ': f'г.  Москва.    		           							 «{day}» {month} 2024 г.  ',
        'Индивидуальный предприниматель Холин Сергей Сергеевич, ОГРНИП 324527500068927, именуемый далее «Агент», с одной стороны, и гражданин РФ _______ паспорт: серия ____ номер ____ , выдан: ___________ , дата выдачи: _________ , адрес регистрации: _____________ , именуемый(ая) далее «Принципал», с другой стороны, а вместе именуемые «Стороны», заключили настоящий Договор о нижеследующем:': f'Индивидуальный предприниматель Холин Сергей Сергеевич, ОГРНИП 324527500068927, именуемый далее «Агент», с одной стороны, и гражданин РФ {name} паспорт: серия {seria} номер {number} , выдан: {kem_vidan} , дата выдачи: {date_passport} , адрес регистрации: {address} , именуемый(ая) далее «Принципал», с другой стороны, а вместе именуемые «Стороны», заключили настоящий Договор о нижеследующем:',
    }


def create_word_file(
        replace_dict: dict,
        list_work: str,
        len_worker: str,
        name: str,
        seria: str,
        number: str,
        kem_vidan: str,
        date_passport: str,
        address: str,
        account_number: str,
        bank_name: str,
        bic: str,
        corr: str,
        telegram_id: int,
):
    doc = Document('/bot/Агентский договор пример.docx')
    for i, paragraph in enumerate(doc.paragraphs):
        # print(f"Абзац {i + 1}: {paragraph.text}")
        for key, value in replace_dict.items():
            paragraph.text = paragraph.text.replace(key, value)

    for table_idx, table in enumerate(doc.tables):
        # print(f"Таблица {table_idx + 1}:")
        for row_idx, row in enumerate(table.rows):
            row_text = [cell.text for cell in row.cells]
            # print(f"  Строка {row_idx + 1}: {', '.join(row_text)}")

    workers_table = doc.tables[0]
    workers_table.cell(1, 0).text = list_work
    workers_table.cell(1, 1).text = len_worker

    data_table = doc.tables[1]
    data_table.cell(0, 0).text = f'''Агент:


Индивидуальный предприниматель Холин Сергей Сергеевич
ИНН: 524613254304
ОГРНИП: 324527500068927
Адрес: Нижегородская обл., г. Бор, МКР. 2-й, д. 46, кв.82.
Расчётный счёт: 40802810620000327428
Название банка:
ООО "Банк Точка"
БИК:
044525104
Корреспондентский счёт:
30101810745374525104

______________/Холин С.С./'''
    data_table.cell(0, 1).text = f'''Принципал:


{name}

Паспорт: серия {seria} номер {number}
Паспорт выдан: {kem_vidan}
Дата выдачи: {date_passport}
Адрес регистрации: {address}
Расчётный счёт: {account_number}
Наименование: {bank_name}
БИК: {bic}
Корреспондентский счёт: {corr}


_______________/{name}

'''

    doc.save(f'/bot/Агентский договор {telegram_id}.docx')

    convertapi.api_secret = settings.SECRET_API
    convertapi.convert('pdf', {
        'File': f'/bot/Агентский договор {telegram_id}.docx'
    }, from_format='docx').save_files(f'/bot/Агентский договор {telegram_id}.pdf')